# TEXT EMBEDDING
from dotenv import load_dotenv
load_dotenv()
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters.markdown import MarkdownHeaderTextSplitter
from langchain_text_splitters.character import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
import numpy as np
import re

doc = DocxDocument(r"D:\PythonFolder\nlp_foundation\Introduction_to_Data_and_Data_Science.docx")
full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
pages = [Document(
    page_content=full_text,
    metadata={"source": "your_file.docx", "page": 1}
)]
md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on = [("#", "course title"),("##", "lecture title")])
pages_md_split = md_splitter.split_text(pages[0].page_content)

for i in range(len(pages_md_split)):
    pages_md_split[i].page_content = ' '.join(pages_md_split[i].page_content.split())

char_splitter = CharacterTextSplitter(
    separator=".",
    chunk_size=500,
    chunk_overlap=50
)

pages_char_split = char_splitter.split_documents(pages_md_split)
print(pages_char_split)

embedding = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# len of all = 1536 -> find similarity - take dot product
vector1 = embedding.embed_query(pages_char_split[3].page_content)
vector2 = embedding.embed_query(pages_char_split[5].page_content)
vector3 = embedding.embed_query(pages_char_split[18].page_content)

print(np.dot(vector1, vector2))
print(np.dot(vector1, vector3))
print(np.dot(vector2, vector3))
print()
print(np.linalg.norm(vector1))
print(np.linalg.norm(vector2))
print(np.linalg.norm(vector3))




#Mark Down Hearder - Based on splitting concerning headers. Different heading levels are represented with various number of hash signs
#Heading1
##Heading2
###Heading3

# from docx import Document as DocxDocument
# from langchain_core.documents import Document
# from langchain_text_splitters.markdown import MarkdownHeaderTextSplitter
# import re


# # Load DOCX
# doc = DocxDocument(r"D:\PythonFolder\nlp_foundation\Introduction_to_Data_and_Data_Science.docx")

# # Extract all paragraph text
# full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])



# # Create LangChain Document (single page)
# pages = [Document(
#     page_content=full_text,
#     metadata={"source": "your_file.docx", "page": 1}
# )]
# # print(pages)

# md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on = [("#", "course title"),("##", "lecture title")])
# pages_md_split = md_splitter.split_text(pages[0].page_content)

# print(pages_md_split)


#Character Text Splitter
# from docx import Document as DocxDocument
# from langchain_core.documents import Document
# import re
# import copy
# from langchain_text_splitters.character import CharacterTextSplitter # segments based on separator, max chunk size, chunk overlap

# def clean_text(text: str) -> str:
#     text = re.sub(r"\s+", " ", text)
#     return text.strip()

# # Load DOCX
# doc = DocxDocument(r"D:\PythonFolder\nlp_foundation\Introduction_to_Data_and_Data_Science.docx")

# # Extract all paragraph text
# full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

# # Clean text
# cleaned_text = clean_text(full_text)

# # Create LangChain Document (single page)
# documents = [Document(
#     page_content=cleaned_text,
#     metadata={"source": "your_file.docx", "page": 1}
# )]

# documents[0].page_content

# #CHunk Overlap: the number of characters overlapping between subsequent chunks (taking some characters from prev chunk-like to build continuity)

# char_splitter = CharacterTextSplitter(separator =".", chunk_size=500, chunk_overlap=0)
# pages_char_split  = char_splitter.split_documents(documents)
# print(len(pages_char_split [1].page_content))
# print(pages_char_split [0].page_content)
# print()
# print(pages_char_split [1].page_content)

